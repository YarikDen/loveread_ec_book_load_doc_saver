#import of required libraries
from bs4 import BeautifulSoup
import requests as req
import docx

doc = docx.Document()
id_book = input("Enter id of book:")
nums = input("Enter number of pages: ")
name_book = input("Input name of book: ")
main_url = "http://loveread.ec/read_book.php?id="+ str(id_book)
page_url = ""
str_text = ""
start_str = 0
end_str = 0
index_header = []
is_par = True
map_text = []
temp_str = ""
resp = ""
net_is_work = True

for num in range(1,int(nums)+1):
        page_url =main_url + "&p=" + str(num)
        #Catch exception if haven`t network connection
        try: 
                resp = req.get(page_url)
        except:
                print("No network error")
                net_is_work = False
                break
        #Make response_context in specified format
        soup = BeautifulSoup(resp.text, 'lxml')
        str_text =  ""
        for tag in soup.find_all("div", attrs={"class":"MsoNormal"}): 
                str_text = (tag.text).replace("\n\n\n\nСтраница\n\n\n\n\n","",2)
                #print(str_text)
        #Find positions of headers
        start_str = 0
        index_header.clear()
        while True:
                start_str = str_text.find("\n\n",start_str) 
                if (start_str != -1):
                        index_header.append(start_str)
                        start_str += 1
                else:
                        break
        ###############################
        #Extract paragraphs and headers
        start_str = 0
        end_str = 0
        is_par = True
        map_text.clear()
        temp_str = ""
        if(str_text[0] != '\n'):
                end_str = str_text.find("\n",start_str)
                temp_str = str_text[start_str:end_str]  
                map_text.append([is_par,temp_str])
                start_str = end_str


        while True:
                #Find start
                start_str = str_text.find("\n",start_str)  
                for a in index_header:
                        #Check for headers
                        if(start_str == a): 
                                is_par = False

                if(is_par != True):
                        end_str = str_text.find("\n",start_str+2)
                else:
                        end_str = str_text.find("\n",start_str+1)

                if(end_str == -1):
                        temp_str = str_text[start_str:]  
                        map_text.append([is_par,temp_str])
                        break

                temp_str = (str_text[start_str+1:end_str]).strip('\n')
                map_text.append([is_par,temp_str])        
                start_str = end_str
                is_par = True
                
        #Load paragraphs and headers to docx-file
        for a in map_text:
                if(a[0] == True):
                        (doc.add_paragraph(a[1])).style = "Normal"
                else:
                        (doc.add_paragraph(a[1])).style = "Heading 1"

if (net_is_work == True):
        #Save doc
        doc.save(name_book+".docx")
        print("Doc was successfully saved")



